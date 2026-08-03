"""
Regression suite for Word tools (M8) - real temp .docx files, no mocking.
"""

import pytest
from docx import Document

from app.core.approval import AutoApproveHandler
from app.core.call_tool import call_tool
from app.core.exceptions import ToolExecutionError


@pytest.fixture
def sample_docx(tmp_path):
    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This mentions bananas specifically.")
    doc.save(file_path)
    return str(file_path)


class TestWordReadTools:
    def test_get_metadata(self, sample_docx, isolated_db):
        result = call_tool("docx_get_metadata", {"file_path": sample_docx})
        assert result.success is True
        assert result.data["paragraph_count"] == 2

    def test_extract_text(self, sample_docx, isolated_db):
        result = call_tool("docx_extract_text", {"file_path": sample_docx})
        assert result.success is True
        assert "bananas" in result.data["paragraphs"][1]


class TestWordWriteTools:
    def test_create_document_requires_approval(self, tmp_path, isolated_db):
        new_path = str(tmp_path / "new.docx")
        with pytest.raises(ToolExecutionError):
            call_tool("docx_create_document", {"file_path": new_path, "initial_text": "hello"})

    def test_create_document_persists(self, tmp_path, isolated_db):
        new_path = str(tmp_path / "new.docx")
        call_tool(
            "docx_create_document",
            {"file_path": new_path, "initial_text": "hello world"},
            approval_handler=AutoApproveHandler(),
        )
        doc = Document(new_path)
        assert doc.paragraphs[0].text == "hello world"

    def test_create_document_refuses_overwrite(self, sample_docx, isolated_db):
        with pytest.raises(ToolExecutionError):
            call_tool(
                "docx_create_document",
                {"file_path": sample_docx, "initial_text": "x"},
                approval_handler=AutoApproveHandler(),
            )

    def test_append_paragraph_persists(self, sample_docx, isolated_db):
        call_tool(
            "docx_append_paragraph",
            {"file_path": sample_docx, "text": "A third paragraph."},
            approval_handler=AutoApproveHandler(),
        )
        doc = Document(sample_docx)
        assert doc.paragraphs[-1].text == "A third paragraph."

    def test_find_and_replace_persists(self, sample_docx, isolated_db):
        result = call_tool(
            "docx_find_and_replace",
            {"file_path": sample_docx, "find_text": "bananas", "replace_text": "oranges"},
            approval_handler=AutoApproveHandler(),
        )
        assert result.data["replacements_made"] == 1
        doc = Document(sample_docx)
        assert "oranges" in doc.paragraphs[1].text
        assert "bananas" not in doc.paragraphs[1].text