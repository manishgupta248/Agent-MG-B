"""
Word READ tools - metadata and text extraction. READ permission.

Note (RAM discipline): python-docx has no read-only/streaming mode -
unlike openpyxl, there's no lazy-loading equivalent for .docx files.
Every open fully parses the document's XML parts. The discipline here
is limited to "don't do unnecessarily expensive extra work" (e.g. a
max_paragraphs cap on extraction) rather than true streaming, which
Word's file format doesn't offer a mechanism for.
"""

from typing import Optional

from docx import Document
from pydantic import BaseModel, Field

from app.core.exceptions import ValidationError
from app.models.tool_result import ToolResult
from app.registry.tool_contract import PermissionLevel, tool
from plugins.word._shared import validate_docx_path

DEFAULT_MAX_PARAGRAPHS = 500


class GetMetadataInput(BaseModel):
    file_path: str = Field(description="Path to the .docx file")


@tool(
    name="docx_get_metadata",
    description="Get a Word document's core properties (title, author, etc.) and paragraph count.",
    permission=PermissionLevel.READ,
    input_schema=GetMetadataInput,
)
def docx_get_metadata(input_data: GetMetadataInput) -> ToolResult:
    path = validate_docx_path(input_data.file_path)
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValidationError(f"Failed to open Word document '{input_data.file_path}': {e}") from e

    props = doc.core_properties
    return ToolResult(success=True, data={
        "paragraph_count": len(doc.paragraphs),
        "title": props.title,
        "author": props.author,
        "subject": props.subject,
        "created": str(props.created) if props.created else None,
    })


class ExtractTextInput(BaseModel):
    file_path: str = Field(description="Path to the .docx file")
    max_paragraphs: int = Field(default=DEFAULT_MAX_PARAGRAPHS, description="Cap on number of paragraphs returned")


@tool(
    name="docx_extract_text",
    description="Extract paragraph text from a Word document, capped at max_paragraphs.",
    permission=PermissionLevel.READ,
    input_schema=ExtractTextInput,
)
def docx_extract_text(input_data: ExtractTextInput) -> ToolResult:
    path = validate_docx_path(input_data.file_path)
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValidationError(f"Failed to open Word document '{input_data.file_path}': {e}") from e

    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if i >= input_data.max_paragraphs:
            break
        paragraphs.append(para.text)

    return ToolResult(success=True, data={
        "paragraph_count_returned": len(paragraphs),
        "paragraphs": paragraphs,
    })