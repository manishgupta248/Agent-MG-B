"""
Word WRITE tools - create_document, append_paragraph, find_and_replace.
All MODIFY permission (approval-gated via call_tool).

docx_create_document creates a NEW file - no atomic-save dance needed
(nothing existing to risk corrupting), but it refuses to overwrite an
existing path, same principle as excel_create_sheet refusing to
overwrite an existing sheet.

docx_append_paragraph and docx_find_and_replace operate on an EXISTING
file and use the same atomic temp-file-then-replace save pattern as
Excel's write tools (plugins/excel/write_tools.py) via
save_docx_atomically in _shared.py.
"""

from pydantic import BaseModel, Field

from app.core.exceptions import ValidationError
from app.models.tool_result import ToolResult
from app.registry.tool_contract import PermissionLevel, tool
from plugins.word._shared import save_docx_atomically, validate_docx_path
from docx import Document


class CreateDocumentInput(BaseModel):
    file_path: str = Field(description="Path where the new .docx file should be created")
    initial_text: str = Field(default="", description="Optional first paragraph's text")


@tool(
    name="docx_create_document",
    description="Create a new Word document. Refuses to overwrite an existing file.",
    permission=PermissionLevel.MODIFY,
    input_schema=CreateDocumentInput,
)
def docx_create_document(input_data: CreateDocumentInput) -> ToolResult:
    path = validate_docx_path(input_data.file_path, must_exist=False)
    if path.exists():
        raise ValidationError(f"File already exists - refusing to overwrite: {input_data.file_path}")

    doc = Document()
    if input_data.initial_text:
        doc.add_paragraph(input_data.initial_text)
    doc.save(str(path))  # brand new file - no existing content to protect, plain save is fine

    return ToolResult(success=True, data={"file_path": input_data.file_path, "created": True})


class AppendParagraphInput(BaseModel):
    file_path: str = Field(description="Path to an existing .docx file")
    text: str = Field(description="Text of the paragraph to append")


@tool(
    name="docx_append_paragraph",
    description="Append a new paragraph to the end of an existing Word document, saving atomically.",
    permission=PermissionLevel.MODIFY,
    input_schema=AppendParagraphInput,
)
def docx_append_paragraph(input_data: AppendParagraphInput) -> ToolResult:
    path = validate_docx_path(input_data.file_path)
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValidationError(f"Failed to open Word document '{input_data.file_path}': {e}") from e

    doc.add_paragraph(input_data.text)
    save_docx_atomically(doc, path)

    return ToolResult(success=True, data={"file_path": input_data.file_path, "appended_text": input_data.text})


class FindAndReplaceInput(BaseModel):
    file_path: str = Field(description="Path to an existing .docx file")
    find_text: str = Field(description="Text to search for")
    replace_text: str = Field(description="Text to replace it with")


@tool(
    name="docx_find_and_replace",
    description=(
        "Find and replace text in a Word document, saving atomically. Operates at the "
        "run level - correctly replaces text fully contained within a single formatting "
        "run, but does NOT handle text that spans across multiple runs (e.g. split by a "
        "bold/italic formatting boundary mid-phrase) - a known, documented limitation."
    ),
    permission=PermissionLevel.MODIFY,
    input_schema=FindAndReplaceInput,
)
def docx_find_and_replace(input_data: FindAndReplaceInput) -> ToolResult:
    path = validate_docx_path(input_data.file_path)
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ValidationError(f"Failed to open Word document '{input_data.file_path}': {e}") from e

    replacements_made = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if input_data.find_text in run.text:
                run.text = run.text.replace(input_data.find_text, input_data.replace_text)
                replacements_made += 1

    save_docx_atomically(doc, path)

    return ToolResult(success=True, data={
        "file_path": input_data.file_path,
        "replacements_made": replacements_made,
    })